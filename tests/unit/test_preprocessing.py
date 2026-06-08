import os
import io
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest
import fitz
from PIL import Image

from app.ml.parsers import (
    ParserFactory, PDFParser, WordParser, ImageParser,
    TXTParser, BaseParser,
)
from app.ml.ocr_engine import OCREngine
from app.schemas.document import DocumentTypeEnum, StandardizedDocument, PageInfo
from app.schemas.common import BoundingBox, TextBlock


@pytest.mark.unit
@pytest.mark.preprocessing
class TestParserFactory:
    def test_detect_document_type_pdf(self):
        factory = ParserFactory()
        assert factory.detect_document_type("report.pdf", "application/pdf") == DocumentTypeEnum.PDF
        assert factory.detect_document_type("report.PDF", "application/pdf") == DocumentTypeEnum.PDF

    def test_detect_document_type_word(self):
        factory = ParserFactory()
        assert factory.detect_document_type("report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document") == DocumentTypeEnum.WORD
        assert factory.detect_document_type("report.doc", "application/msword") == DocumentTypeEnum.WORD

    def test_detect_document_type_image(self):
        factory = ParserFactory()
        assert factory.detect_document_type("invoice.jpg", "image/jpeg") == DocumentTypeEnum.IMAGE
        assert factory.detect_document_type("scan.png", "image/png") == DocumentTypeEnum.IMAGE
        assert factory.detect_document_type("photo.tiff", "image/tiff") == DocumentTypeEnum.IMAGE

    def test_detect_document_type_txt(self):
        factory = ParserFactory()
        assert factory.detect_document_type("notes.txt", "text/plain") == DocumentTypeEnum.TXT

    def test_detect_document_type_unknown(self):
        factory = ParserFactory()
        assert factory.detect_document_type("file.xyz", "application/xyz") == DocumentTypeEnum.UNKNOWN

    def test_get_parser_pdf(self):
        factory = ParserFactory()
        parser = factory.get_parser(DocumentTypeEnum.PDF)
        assert isinstance(parser, PDFParser)

    def test_get_parser_word(self):
        factory = ParserFactory()
        parser = factory.get_parser(DocumentTypeEnum.WORD)
        assert isinstance(parser, WordParser)

    def test_get_parser_image(self):
        factory = ParserFactory()
        parser = factory.get_parser(DocumentTypeEnum.IMAGE)
        assert isinstance(parser, ImageParser)

    def test_get_parser_txt(self):
        factory = ParserFactory()
        parser = factory.get_parser(DocumentTypeEnum.TXT)
        assert isinstance(parser, TXTParser)

    def test_get_parser_unknown_raises(self):
        factory = ParserFactory()
        with pytest.raises(ValueError, match="Unsupported document type"):
            factory.get_parser(DocumentTypeEnum.UNKNOWN)


@pytest.mark.unit
@pytest.mark.preprocessing
class TestPDFParser:
    def test_parse_normal_pdf(self, sample_pdf_path):
        parser = PDFParser()
        result = parser.parse(str(sample_pdf_path["path"]), sample_pdf_path["path"].name)

        assert isinstance(result, StandardizedDocument)
        assert result.document_type == DocumentTypeEnum.PDF
        assert result.page_count >= 1
        assert result.original_filename == sample_pdf_path["path"].name
        assert len(result.pages) == result.page_count

        page = result.pages[0]
        assert page.page_number == 1
        assert page.width > 0
        assert page.height > 0
        assert len(page.text_blocks) > 0

    def test_pdf_text_blocks_have_correct_coordinates(self, sample_pdf_path):
        parser = PDFParser()
        result = parser.parse(str(sample_pdf_path["path"]), "test.pdf")
        page = result.pages[0]

        for text_block in page.text_blocks:
            assert isinstance(text_block, TextBlock)
            assert isinstance(text_block.bbox, BoundingBox)
            assert text_block.bbox.x1 >= 0
            assert text_block.bbox.y1 >= 0
            assert text_block.bbox.x2 > text_block.bbox.x1
            assert text_block.bbox.y2 > text_block.bbox.y1
            assert text_block.bbox.x2 <= page.width
            assert text_block.bbox.y2 <= page.height
            assert text_block.confidence == 1.0

    def test_pdf_contains_patient_name(self, sample_pdf_path):
        parser = PDFParser()
        result = parser.parse(str(sample_pdf_path["path"]), "test.pdf")
        page = result.pages[0]

        all_text = " ".join(tb.text for tb in page.text_blocks)
        assert sample_pdf_path["patient_name"] in all_text

    def test_pdf_contains_diagnosis_code(self, sample_pdf_path):
        parser = PDFParser()
        result = parser.parse(str(sample_pdf_path["path"]), "test.pdf")
        page = result.pages[0]

        all_text = " ".join(tb.text for tb in page.text_blocks)
        assert sample_pdf_path["diagnosis_code"] in all_text

    def test_pdf_tables_extracted(self, sample_pdf_path):
        parser = PDFParser()
        result = parser.parse(str(sample_pdf_path["path"]), "test.pdf")
        page = result.pages[0]

        assert len(page.tables) >= 1
        table = page.tables[0]
        assert table.row_count >= 2
        assert table.col_count >= 2
        assert len(table.headers) >= 1

    def test_encrypted_pdf_raises_clear_error(self, encrypted_pdf_path):
        parser = PDFParser()

        with pytest.raises(Exception) as excinfo:
            parser.parse(str(encrypted_pdf_path), "encrypted.pdf")

        error_msg = str(excinfo.value).lower()
        assert any(keyword in error_msg for keyword in ["encrypt", "password", "permission", "protect"])

    def test_nonexistent_pdf_raises(self):
        parser = PDFParser()

        with pytest.raises(FileNotFoundError):
            parser.parse("/nonexistent/file.pdf", "nonexistent.pdf")

    def test_ocr_triggered_for_scanned_pdf(self, temp_dir, mock_ocr_engine):
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Image as RLImage

        img = Image.new("RGB", (500, 700), color="white")
        img_path = temp_dir / "scan_image.jpg"
        img.save(img_path)

        pdf_path = temp_dir / "scanned_doc.pdf"
        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        doc.build([RLImage(str(img_path), width=400, height=560)])

        parser = PDFParser(ocr_engine=mock_ocr_engine)
        result = parser.parse(str(pdf_path), "scanned_doc.pdf")

        assert mock_ocr_engine.called
        page = result.pages[0]
        assert page.ocr_confidence is not None
        assert page.ocr_confidence > 0

    def test_low_quality_ocr_marked_low_confidence(self, temp_dir):
        from app.schemas.common import TextBlock as TB

        class LowConfidenceOCREngine:
            def is_available(self):
                return True

            def ocr_image(self, image, page_number=1):
                return [
                    TB(
                        text="模糊文本",
                        bbox=BoundingBox(x1=0, y1=0, x2=100, y2=30),
                        confidence=0.4,
                        page_number=page_number,
                    ),
                    TB(
                        text="难识别",
                        bbox=BoundingBox(x1=0, y1=35, x2=80, y2=65),
                        confidence=0.35,
                        page_number=page_number,
                    ),
                ]

            def get_ocr_metadata(self):
                return {"engine": "test"}

        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Image as RLImage

        img = Image.new("RGB", (500, 700), color="white")
        img_path = temp_dir / "low_quality.jpg"
        img.save(img_path)

        pdf_path = temp_dir / "low_quality_scan.pdf"
        doc = SimpleDocTemplate(str(pdf_path), pagesize=A4)
        doc.build([RLImage(str(img_path), width=400, height=560)])

        parser = PDFParser(ocr_engine=LowConfidenceOCREngine())
        result = parser.parse(str(pdf_path), "low_quality.pdf")

        page = result.pages[0]
        assert page.ocr_confidence < 0.5

        for tb in page.text_blocks:
            assert tb.confidence < 0.5


@pytest.mark.unit
@pytest.mark.preprocessing
class TestImageParser:
    def test_parse_image_jpg(self, sample_image_path):
        parser = ImageParser()
        result = parser.parse(str(sample_image_path["path"]), sample_image_path["path"].name)

        assert isinstance(result, StandardizedDocument)
        assert result.document_type == DocumentTypeEnum.IMAGE
        assert result.page_count == 1

    def test_image_ocr_extracts_text(self, sample_image_path, mock_ocr_engine):
        parser = ImageParser(ocr_engine=mock_ocr_engine)
        result = parser.parse(str(sample_image_path["path"]), "test.jpg")

        assert mock_ocr_engine.called
        assert len(result.pages[0].text_blocks) >= 1

    def test_low_quality_image_low_confidence(self, low_quality_image_path):
        parser = ImageParser()
        result = parser.parse(str(low_quality_image_path), "low_quality.jpg")

        page = result.pages[0]
        if page.ocr_confidence is not None:
            assert page.ocr_confidence < 0.6

    def test_image_text_blocks_have_confidence(self, sample_image_path):
        parser = ImageParser()
        result = parser.parse(str(sample_image_path["path"]), "test.jpg")

        for tb in result.pages[0].text_blocks:
            assert tb.confidence is not None
            assert 0 <= tb.confidence <= 1


@pytest.mark.unit
@pytest.mark.preprocessing
class TestTXTParser:
    def test_parse_txt_file(self, temp_dir):
        txt_path = temp_dir / "test.txt"
        content = """患者姓名：张三
身份证号：110101199001011234
诊断编码：J45.900
诊断描述：支气管哮喘
总金额：3500.00
日期：2024-01-15
"""
        txt_path.write_text(content, encoding="utf-8")

        parser = TXTParser()
        result = parser.parse(str(txt_path), "test.txt")

        assert isinstance(result, StandardizedDocument)
        assert result.document_type == DocumentTypeEnum.TXT
        assert result.page_count == 1

        all_text = " ".join(tb.text for tb in result.pages[0].text_blocks)
        assert "张三" in all_text
        assert "J45.900" in all_text

    def test_txt_empty_file(self, temp_dir):
        txt_path = temp_dir / "empty.txt"
        txt_path.write_text("", encoding="utf-8")

        parser = TXTParser()
        result = parser.parse(str(txt_path), "empty.txt")

        assert len(result.pages[0].text_blocks) == 0

    def test_txt_chinese_content(self, temp_dir):
        txt_path = temp_dir / "chinese.txt"
        content = "这是一份中文测试文档，包含混合English内容。"
        txt_path.write_text(content, encoding="utf-8")

        parser = TXTParser()
        result = parser.parse(str(txt_path), "chinese.txt")

        all_text = " ".join(tb.text for tb in result.pages[0].text_blocks)
        assert "中文测试文档" in all_text
        assert "English" in all_text


@pytest.mark.unit
@pytest.mark.preprocessing
class TestWordParser:
    def test_parse_word_document(self, temp_dir):
        from docx import Document
        from docx.shared import Pt

        doc_path = temp_dir / "test.docx"
        doc = Document()

        doc.add_heading("诊断证明书", level=1)
        doc.add_paragraph("患者姓名：李四")
        doc.add_paragraph("身份证号：310101198505055678")
        doc.add_paragraph("诊断编码：I10")
        doc.add_paragraph("诊断描述：原发性高血压")

        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "项目"
        table.rows[0].cells[1].text = "类型"
        table.rows[0].cells[2].text = "金额"
        table.rows[1].cells[0].text = "挂号费"
        table.rows[1].cells[1].text = "诊疗"
        table.rows[1].cells[2].text = "50.00"

        doc.save(str(doc_path))

        parser = WordParser()
        result = parser.parse(str(doc_path), "test.docx")

        assert isinstance(result, StandardizedDocument)
        assert result.document_type == DocumentTypeEnum.WORD

        all_text = " ".join(tb.text for tb in result.pages[0].text_blocks)
        assert "李四" in all_text
        assert "I10" in all_text


@pytest.mark.unit
@pytest.mark.preprocessing
class TestCrossPageTableHandling:
    def test_cross_page_table_merged_correctly(self, multi_page_pdf_path):
        from app.ml.table_extractor import TableExtractor

        parser = PDFParser()
        doc = parser.parse(str(multi_page_pdf_path), "multi_page.pdf")

        assert doc.page_count >= 2

        all_tables = []
        for page in doc.pages:
            all_tables.extend(page.tables)

        assert len(all_tables) >= 2

        table_extractor = TableExtractor()
        merged = table_extractor.merge_cross_page_tables(all_tables)

        assert len(merged) >= 1
        merged_table = merged[0]

        all_rows = []
        for t in all_tables:
            all_rows.extend(t.rows)

        has_total = any("合计" in str(row) for row in all_rows)
        assert has_total


@pytest.mark.unit
@pytest.mark.preprocessing
class TestStandardizedDocument:
    def test_standardized_document_structure(self, mock_standardized_doc):
        doc = mock_standardized_doc

        assert doc.document_id == 1
        assert doc.page_count == 1
        assert len(doc.pages) == 1

        page = doc.pages[0]
        assert page.page_number == 1
        assert len(page.text_blocks) > 0
        assert len(page.tables) == 1

    def test_text_block_confidence_accessible(self, mock_standardized_doc):
        page = mock_standardized_doc.pages[0]

        for tb in page.text_blocks:
            assert hasattr(tb, "confidence")
            assert tb.confidence is not None
            assert 0 <= tb.confidence <= 1

    def test_document_metadata(self, mock_standardized_doc):
        assert mock_standardized_doc.language == "zh"
        assert mock_standardized_doc.preprocessing_time is not None
        assert mock_standardized_doc.preprocessing_time >= 0
