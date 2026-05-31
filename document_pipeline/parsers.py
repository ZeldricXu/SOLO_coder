from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional
import io
import os
from pathlib import Path

from .schemas import DocumentFormat, DocumentInfo
from common.logger import get_logger

logger = get_logger(__name__)


class BaseParser(ABC):
    format: DocumentFormat

    @abstractmethod
    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        pass


class PDFParser(BaseParser):
    format = DocumentFormat.PDF

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"pages": 0, "images": [], "tables": []}
        try:
            import pdfplumber
            if file_path:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    metadata["pages"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
                        if kwargs.get("extract_tables"):
                            tables = page.extract_tables()
                            if tables:
                                metadata["tables"].extend(tables)
                    return text, metadata
            elif file_content:
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    text = ""
                    metadata["pages"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        text += page_text + "\n\n"
                    return text, metadata
        except ImportError:
            logger.warning("pdfplumber not installed, using fallback")
            return "PDF content (parsing requires pdfplumber)", metadata
        except Exception as e:
            logger.error(f"PDF parsing error: {str(e)}")
            return "", metadata


class DOCXParser(BaseParser):
    format = DocumentFormat.DOCX

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"paragraphs": 0, "tables": []}
        try:
            from docx import Document
            if file_path:
                doc = Document(file_path)
            elif file_content:
                doc = Document(io.BytesIO(file_content))
            else:
                return "", metadata

            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            metadata["paragraphs"] = len(doc.paragraphs)

            if kwargs.get("extract_tables"):
                for table in doc.tables:
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    metadata["tables"].append(table_data)

            return text, metadata
        except ImportError:
            logger.warning("python-docx not installed")
            return "DOCX content (parsing requires python-docx)", metadata
        except Exception as e:
            logger.error(f"DOCX parsing error: {str(e)}")
            return "", metadata


class TXTParser(BaseParser):
    format = DocumentFormat.TXT

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"encoding": "utf-8", "lines": 0}
        try:
            if file_path:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
            elif file_content:
                content = file_content.decode("utf-8", errors="ignore")
            else:
                return "", metadata

            metadata["lines"] = content.count("\n") + 1
            return content, metadata
        except Exception as e:
            logger.error(f"TXT parsing error: {str(e)}")
            return "", metadata


class MDParser(BaseParser):
    format = DocumentFormat.MD

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"headings": [], "code_blocks": 0}
        try:
            if file_path:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
            elif file_content:
                content = file_content.decode("utf-8", errors="ignore")
            else:
                return "", metadata

            for line in content.split("\n"):
                if line.startswith("#"):
                    metadata["headings"].append(line.strip())
                if line.startswith("```"):
                    metadata["code_blocks"] += 1

            return content, metadata
        except Exception as e:
            logger.error(f"MD parsing error: {str(e)}")
            return "", metadata


class HTMLParser(BaseParser):
    format = DocumentFormat.HTML

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"links": [], "images": []}
        try:
            from bs4 import BeautifulSoup
            if file_path:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()
            elif file_content:
                content = file_content.decode("utf-8", errors="ignore")
            else:
                return "", metadata

            soup = BeautifulSoup(content, "html.parser")

            for tag in soup(["script", "style"]):
                tag.decompose()

            text = soup.get_text(separator="\n", strip=True)

            for link in soup.find_all("a"):
                href = link.get("href")
                if href:
                    metadata["links"].append(href)

            for img in soup.find_all("img"):
                src = img.get("src")
                if src:
                    metadata["images"].append(src)

            return text, metadata
        except ImportError:
            logger.warning("beautifulsoup4 not installed")
            return "HTML content (parsing requires beautifulsoup4)", metadata
        except Exception as e:
            logger.error(f"HTML parsing error: {str(e)}")
            return "", metadata


class CSVParser(BaseParser):
    format = DocumentFormat.CSV

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"rows": 0, "columns": []}
        try:
            import csv
            if file_path:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    reader = csv.reader(f)
                    rows = list(reader)
            elif file_content:
                content = file_content.decode("utf-8", errors="ignore")
                reader = csv.reader(io.StringIO(content))
                rows = list(reader)
            else:
                return "", metadata

            if rows:
                metadata["columns"] = rows[0]
                metadata["rows"] = len(rows) - 1

            text = "\n".join([",".join(row) for row in rows])
            return text, metadata
        except Exception as e:
            logger.error(f"CSV parsing error: {str(e)}")
            return "", metadata


class XLSXParser(BaseParser):
    format = DocumentFormat.XLSX

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        metadata: Dict[str, Any] = {"sheets": [], "total_rows": 0}
        try:
            from openpyxl import load_workbook
            if file_path:
                wb = load_workbook(file_path, read_only=True, data_only=True)
            elif file_content:
                wb = load_workbook(io.BytesIO(file_content), read_only=True, data_only=True)
            else:
                return "", metadata

            text_parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_info = {"name": sheet_name, "rows": ws.max_row, "columns": ws.max_column}
                metadata["sheets"].append(sheet_info)
                metadata["total_rows"] += ws.max_row

                for row in ws.iter_rows(values_only=True):
                    text_parts.append(",".join([str(cell) if cell is not None else "" for cell in row]))

            return "\n".join(text_parts), metadata
        except ImportError:
            logger.warning("openpyxl not installed")
            return "XLSX content (parsing requires openpyxl)", metadata
        except Exception as e:
            logger.error(f"XLSX parsing error: {str(e)}")
            return "", metadata


class JSONParser(BaseParser):
    format = DocumentFormat.JSON

    def parse(self, file_path: Optional[str] = None, file_content: Optional[bytes] = None, **kwargs) -> tuple[str, Dict[str, Any]]:
        import json
        metadata: Dict[str, Any] = {"keys": [], "is_array": False}
        try:
            if file_path:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    data = json.load(f)
            elif file_content:
                data = json.loads(file_content.decode("utf-8", errors="ignore"))
            else:
                return "", metadata

            metadata["is_array"] = isinstance(data, list)

            def extract_keys(obj, prefix=""):
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        full_key = f"{prefix}.{k}" if prefix else k
                        metadata["keys"].append(full_key)
                        extract_keys(v, full_key)
                elif isinstance(obj, list) and obj:
                    extract_keys(obj[0], f"{prefix}[0]")

            extract_keys(data)
            return json.dumps(data, ensure_ascii=False, indent=2), metadata
        except Exception as e:
            logger.error(f"JSON parsing error: {str(e)}")
            return "", metadata


PARSER_MAP: Dict[DocumentFormat, type[BaseParser]] = {
    DocumentFormat.PDF: PDFParser,
    DocumentFormat.DOCX: DOCXParser,
    DocumentFormat.TXT: TXTParser,
    DocumentFormat.MD: MDParser,
    DocumentFormat.HTML: HTMLParser,
    DocumentFormat.CSV: CSVParser,
    DocumentFormat.XLSX: XLSXParser,
    DocumentFormat.JSON: JSONParser,
}


def get_parser(doc_format: DocumentFormat) -> BaseParser:
    parser_class = PARSER_MAP.get(doc_format)
    if not parser_class:
        raise ValueError(f"No parser available for format: {doc_format}")
    return parser_class()


def detect_format(file_name: str) -> Optional[DocumentFormat]:
    ext = Path(file_name).suffix.lower().lstrip(".")
    format_map = {
        "pdf": DocumentFormat.PDF,
        "docx": DocumentFormat.DOCX,
        "doc": DocumentFormat.DOCX,
        "txt": DocumentFormat.TXT,
        "md": DocumentFormat.MD,
        "markdown": DocumentFormat.MD,
        "html": DocumentFormat.HTML,
        "htm": DocumentFormat.HTML,
        "csv": DocumentFormat.CSV,
        "xlsx": DocumentFormat.XLSX,
        "xls": DocumentFormat.XLSX,
        "json": DocumentFormat.JSON,
        "pptx": DocumentFormat.PPTX,
        "eml": DocumentFormat.EML,
        "msg": DocumentFormat.MSG,
    }
    return format_map.get(ext)
