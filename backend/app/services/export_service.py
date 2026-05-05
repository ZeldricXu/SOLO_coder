import os
import tempfile
import shutil
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from zipfile import ZipFile
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.colors import HexColor
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from app.services.excel_chart_generator import ExcelChartGenerator, ChartConfig, ChartType
    HAS_EXCEL_CHART = True
except ImportError:
    HAS_EXCEL_CHART = False

from app.models import Report, ReportSection, survey_store

class ExportService:
    def __init__(self, export_folder: str):
        self.export_folder = export_folder
        os.makedirs(self.export_folder, exist_ok=True)
    
    def export_to_word(self, report_id: str) -> Optional[str]:
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word export")
        
        report = survey_store.get_report(report_id)
        if not report:
            return None
        
        file_path = os.path.join(self.export_folder, f"{report_id}.docx")
        
        doc = Document()
        
        title = doc.add_heading(report.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"生成时间：{report.created_at}")
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.italic = True
        
        doc.add_paragraph()
        doc.add_heading("目录", level=1)
        
        for i, section in enumerate(report.sections):
            toc_para = doc.add_paragraph()
            toc_para.add_run(f"{i+1}. {section.title}")
        
        doc.add_page_break()
        
        for section in report.sections:
            self._add_section_to_word(doc, section)
        
        doc.save(file_path)
        
        return file_path
    
    def _add_section_to_word(self, doc: 'Document', section: ReportSection) -> None:
        heading = doc.add_heading(section.title, level=1)
        heading.paragraph_format.space_before = Pt(24)
        
        content_para = doc.add_paragraph()
        lines = section.content.split('\n')
        for line in lines:
            if line.strip():
                run = content_para.add_run(line + '\n')
                run.font.size = Pt(11)
        
        if section.data:
            doc.add_heading("详细数据", level=2)
            
            if section.section_type == "frequency" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "")
                    data = stat.get("data", {})
                    frequencies = data.get("frequencies", [])
                    
                    doc.add_paragraph()
                    doc.add_heading(field_name, level=3)
                    
                    if frequencies:
                        table = doc.add_table(rows=1, cols=3)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = '选项'
                        hdr_cells[1].text = '数量'
                        hdr_cells[2].text = '百分比'
                        
                        for item in frequencies[:10]:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(item.get('value', ''))
                            row_cells[1].text = str(item.get('count', 0))
                            row_cells[2].text = f"{item.get('percentage', 0)}%"
            
            elif section.section_type == "descriptive" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "")
                    data = stat.get("data", {})
                    
                    doc.add_paragraph()
                    doc.add_heading(field_name, level=3)
                    
                    table = doc.add_table(rows=1, cols=8)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '样本数'
                    hdr_cells[1].text = '均值'
                    hdr_cells[2].text = '中位数'
                    hdr_cells[3].text = '标准差'
                    hdr_cells[4].text = '最小值'
                    hdr_cells[5].text = '最大值'
                    hdr_cells[6].text = '25分位'
                    hdr_cells[7].text = '75分位'
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(data.get('count', 0))
                    row_cells[1].text = f"{data.get('mean', 0):.4f}"
                    row_cells[2].text = f"{data.get('median', 0):.4f}"
                    row_cells[3].text = f"{data.get('std', 0):.4f}"
                    row_cells[4].text = f"{data.get('min', 0):.4f}"
                    row_cells[5].text = f"{data.get('max', 0):.4f}"
                    row_cells[6].text = f"{data.get('q25', 0):.4f}"
                    row_cells[7].text = f"{data.get('q75', 0):.4f}"
        
        doc.add_paragraph()
    
    def export_to_pdf(self, report_id: str) -> Optional[str]:
        if not HAS_REPORTLAB:
            raise ImportError("reportlab is required for PDF export")
        
        report = survey_store.get_report(report_id)
        if not report:
            return None
        
        file_path = os.path.join(self.export_folder, f"{report_id}.pdf")
        
        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=1
        )
        
        heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=styles['Heading1'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=10
        )
        
        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=11,
            leading=18,
            spaceBefore=6,
            spaceAfter=6
        )
        
        story = []
        
        story.append(Paragraph(report.title, title_style))
        story.append(Spacer(1, 12))
        
        subtitle = f"生成时间：{report.created_at}"
        subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=styles['Italic'],
            fontSize=10,
            alignment=1
        )
        story.append(Paragraph(subtitle, subtitle_style))
        story.append(Spacer(1, 30))
        
        story.append(Paragraph("目录", heading1_style))
        story.append(Spacer(1, 12))
        
        for i, section in enumerate(report.sections):
            toc_item = f"{i+1}. {section.title}"
            story.append(Paragraph(toc_item, body_style))
        
        story.append(Spacer(1, 30))
        
        for section in report.sections:
            story.extend(self._build_pdf_section(section, heading1_style, heading2_style, body_style))
        
        doc.build(story)
        
        return file_path
    
    def _build_pdf_section(
        self, 
        section: ReportSection, 
        heading1_style: 'ParagraphStyle',
        heading2_style: 'ParagraphStyle',
        body_style: 'ParagraphStyle'
    ) -> List[Any]:
        story = []
        
        story.append(Paragraph(section.title, heading1_style))
        story.append(Spacer(1, 6))
        
        lines = section.content.split('\n')
        for line in lines:
            if line.strip():
                story.append(Paragraph(line.replace('\n', '<br/>'), body_style))
        
        if section.data:
            story.append(Spacer(1, 12))
            story.append(Paragraph("详细数据", heading2_style))
            story.append(Spacer(1, 6))
            
            if section.section_type == "frequency" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "")
                    data = stat.get("data", {})
                    frequencies = data.get("frequencies", [])
                    
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(field_name, heading2_style))
                    story.append(Spacer(1, 6))
                    
                    if frequencies:
                        table_data = [['选项', '数量', '百分比']]
                        for item in frequencies[:10]:
                            table_data.append([
                                str(item.get('value', '')),
                                str(item.get('count', 0)),
                                f"{item.get('percentage', 0)}%"
                            ])
                        
                        table = Table(table_data, colWidths=[2*inch, 1*inch, 1*inch])
                        table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#4A90E2')),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        story.append(table)
            
            elif section.section_type == "descriptive" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "")
                    data = stat.get("data", {})
                    
                    story.append(Spacer(1, 6))
                    story.append(Paragraph(field_name, heading2_style))
                    story.append(Spacer(1, 6))
                    
                    table_data = [
                        ['样本数', '均值', '中位数', '标准差', '最小值', '最大值'],
                        [
                            str(data.get('count', 0)),
                            f"{data.get('mean', 0):.4f}",
                            f"{data.get('median', 0):.4f}",
                            f"{data.get('std', 0):.4f}",
                            f"{data.get('min', 0):.4f}",
                            f"{data.get('max', 0):.4f}"
                        ]
                    ]
                    
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#4A90E2')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.whitesmoke),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(table)
        
        story.append(Spacer(1, 20))
        
        return story
    
    def get_export_file_path(self, report_id: str, format: str) -> str:
        if format == 'pdf':
            return os.path.join(self.export_folder, f"{report_id}.pdf")
        else:
            return os.path.join(self.export_folder, f"{report_id}.docx")
    
    def _create_excel_chart_from_section(
        self, 
        section: ReportSection,
        chart_generator: ExcelChartGenerator
    ) -> List[Tuple[str, str]]:
        """
        从报告章节创建Excel图表文件
        
        Returns:
            List[(图表标题, Excel文件路径)]
        """
        excel_files = []
        
        if not section.data:
            return excel_files
        
        try:
            if section.section_type == "frequency" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "未知字段")
                    data = stat.get("data", {})
                    
                    chart_config = chart_generator.create_chart_from_analysis_result(
                        "frequency", {**data, "field_name": field_name}
                    )
                    
                    safe_filename = re.sub(r'[^\w\-_\u4e00-\u9fa5]', '_', field_name)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f"频率图表_{safe_filename}_{timestamp}.xlsx"
                    
                    excel_path = chart_generator.generate_excel_chart(chart_config, filename)
                    excel_files.append((field_name, excel_path))
            
            elif section.section_type == "descriptive" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "未知字段")
                    data = stat.get("data", {})
                    
                    chart_config = chart_generator.create_chart_from_analysis_result(
                        "descriptive", {**data, "field_name": field_name}
                    )
                    
                    safe_filename = re.sub(r'[^\w\-_\u4e00-\u9fa5]', '_', field_name)
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f"描述统计图表_{safe_filename}_{timestamp}.xlsx"
                    
                    excel_path = chart_generator.generate_excel_chart(chart_config, filename)
                    excel_files.append((field_name, excel_path))
            
            elif section.section_type == "cross" and "cross_analyses" in section.data:
                for idx, analysis in enumerate(section.data["cross_analyses"]):
                    variables = analysis.get("variables", ["变量1", "变量2"])
                    chart_title = f"交叉分析_{variables[0]}_vs_{variables[1]}"
                    
                    chart_config = chart_generator.create_chart_from_analysis_result(
                        "cross", analysis
                    )
                    
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    filename = f"交叉分析图表_{idx + 1}_{timestamp}.xlsx"
                    
                    excel_path = chart_generator.generate_excel_chart(chart_config, filename)
                    excel_files.append((chart_title, excel_path))
        
        except Exception as e:
            pass
        
        return excel_files
    
    def _format_table_headers(self, table: 'Table') -> None:
        """格式化表格表头"""
        for cell in table.rows[0].cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '4472C4')
            cell._tc.get_or_add_tcPr().append(shading_elm)
            
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
    
    def _add_significance_info_to_word(
        self, 
        doc: 'Document', 
        significance: Dict[str, Any]
    ) -> None:
        """向Word文档添加显著性检验信息"""
        if not significance:
            return
        
        test_type = significance.get("test_type", "未知检验")
        p_value = significance.get("p_value", 0)
        is_significant = significance.get("significant", False)
        details = significance.get("details", {})
        
        doc.add_heading("显著性检验结果", level=4)
        
        test_type_names = {
            "chi_square": "卡方检验",
            "fisher_exact": "Fisher精确检验",
            "t_test": "独立样本t检验",
            "anova": "单因素方差分析",
            "mann_whitney_u": "Mann-Whitney U检验（非参数）",
            "kruskal_wallis": "Kruskal-Wallis H检验（非参数）",
            "wilcoxon": "Wilcoxon符号秩检验"
        }
        
        info_para = doc.add_paragraph()
        info_para.add_run(f"检验方法：{test_type_names.get(test_type, test_type)}\n").bold = True
        info_para.add_run(f"P值：{p_value:.6f}\n")
        
        significance_text = "显著（P < 0.05）" if is_significant else "不显著（P >= 0.05）"
        significance_run = info_para.add_run(f"检验结果：{significance_text}\n")
        significance_run.bold = True
        significance_run.font.color.rgb = RGBColor(0, 128, 0) if is_significant else RGBColor(128, 0, 0)
        
        if details:
            warnings = details.get("warnings", [])
            if warnings:
                warning_para = doc.add_paragraph()
                warning_para.add_run("⚠️ 警告信息：\n").bold = True
                warning_para.add_run(f"检验方法：{test_type_names.get(test_type, test_type)}\n").bold = True
                for warning in warnings:
                    warning_para.add_run(f"  • {warning}\n")
                warning_para.runs[0].font.color.rgb = RGBColor(255, 128, 0)
            
            if details.get("switched_from_chi2"):
                switch_para = doc.add_paragraph()
                switch_para.add_run("📊 检验方法切换说明：\n").bold = True
                switch_para.add_run("  由于期望频数不满足卡方检验假设条件，已自动切换为Fisher精确检验。\n")
            
            if details.get("switched_from_t_test"):
                switch_para = doc.add_paragraph()
                switch_para.add_run("📊 检验方法切换说明：\n").bold = True
                switch_para.add_run("  由于样本量较小，已自动切换为非参数Mann-Whitney U检验。\n")
            
            if details.get("switched_from_anova"):
                switch_para = doc.add_paragraph()
                switch_para.add_run("📊 检验方法切换说明：\n").bold = True
                switch_para.add_run("  由于存在小样本组，已自动切换为非参数Kruskal-Wallis H检验。\n")
            
            if "expected_frequencies" in details:
                exp_freq_para = doc.add_paragraph()
                exp_freq_para.add_run("期望频数信息：\n").bold = True
                exp_freq_para.add_run(f"  最小期望频数：{details.get('min_expected', 'N/A')}\n")
                if details.get("count_below_5"):
                    exp_freq_para.add_run(f"  期望频数 < 5 的单元格数：{details.get('count_below_5')}\n")
                if details.get("count_below_1"):
                    exp_freq_para.add_run(f"  期望频数 < 1 的单元格数：{details.get('count_below_1')}\n")
    
    def export_to_word_with_charts(
        self, 
        report_id: str,
        include_excel_charts: bool = True
    ) -> Tuple[Optional[str], List[str]]:
        """
        导出Word文档，并可选地生成配套的可编辑Excel图表文件
        
        Args:
            report_id: 报告ID
            include_excel_charts: 是否包含Excel图表文件
            
        Returns:
            (Word文件路径, Excel图表文件路径列表)
        """
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word export")
        
        report = survey_store.get_report(report_id)
        if not report:
            return None, []
        
        file_path = os.path.join(self.export_folder, f"{report_id}.docx")
        excel_chart_files = []
        
        chart_generator = None
        if include_excel_charts and HAS_EXCEL_CHART:
            charts_folder = os.path.join(self.export_folder, f"{report_id}_charts")
            os.makedirs(charts_folder, exist_ok=True)
            chart_generator = ExcelChartGenerator(charts_folder)
        
        doc = Document()
        
        title = doc.add_heading(report.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"生成时间：{report.created_at}")
        subtitle_run.font.size = Pt(10)
        subtitle_run.font.italic = True
        
        if include_excel_charts and chart_generator:
            chart_note = doc.add_paragraph()
            chart_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = chart_note.add_run("📊 本报告包含可编辑的Excel图表文件（随报告一同导出）")
            note_run.font.size = Pt(9)
            note_run.font.color.rgb = RGBColor(68, 114, 196)
        
        doc.add_paragraph()
        doc.add_heading("目录", level=1)
        
        for i, section in enumerate(report.sections):
            toc_para = doc.add_paragraph()
            toc_para.add_run(f"{i+1}. {section.title}")
        
        doc.add_page_break()
        
        for section_idx, section in enumerate(report.sections):
            self._add_section_to_word_enhanced(doc, section, section_idx + 1)
            
            if include_excel_charts and chart_generator:
                section_charts = self._create_excel_chart_from_section(section, chart_generator)
                for chart_title, excel_path in section_charts:
                    excel_chart_files.append(excel_path)
                    
                    chart_ref_para = doc.add_paragraph()
                    chart_ref_para.add_run(f"📈 相关图表：{chart_title}").italic = True
                    chart_ref_para.add_run(f"（详见附件Excel文件）").font.size = Pt(9)
        
        doc.add_page_break()
        doc.add_heading("附录：显著性检验说明", level=1)
        
        appendix_para = doc.add_paragraph()
        appendix_para.add_run("检验方法说明：\n\n").bold = True
        
        test_descriptions = [
            ("卡方检验 (Chi-square)", "用于两个分类变量的独立性检验。假设条件：所有期望频数 >= 1，且至少80%的期望频数 >= 5。"),
            ("Fisher精确检验 (Fisher's Exact)", "用于2x2列联表的精确检验。当期望频数不满足卡方检验条件时自动使用。"),
            ("独立样本t检验 (t-test)", "用于比较两个独立样本的均值差异。假设条件：正态分布、方差齐性。"),
            ("Mann-Whitney U检验", "非参数检验，用于比较两个独立样本的分布差异。当样本量较小时自动使用。"),
            ("单因素方差分析 (ANOVA)", "用于比较三个或更多独立样本的均值差异。"),
            ("Kruskal-Wallis H检验", "非参数检验，用于比较三个或更多独立样本的分布差异。"),
        ]
        
        for test_name, description in test_descriptions:
            appendix_para.add_run(f"• {test_name}\n").bold = True
            appendix_para.add_run(f"  {description}\n\n")
        
        appendix_para.add_run("\n显著性水平说明：\n\n").bold = True
        appendix_para.add_run("• P < 0.05：在95%置信水平下，差异具有统计学意义\n")
        appendix_para.add_run("• P < 0.01：在99%置信水平下，差异具有统计学意义\n")
        appendix_para.add_run("• P >= 0.05：差异不具有统计学意义\n")
        
        doc.save(file_path)
        
        return file_path, excel_chart_files
    
    def _add_section_to_word_enhanced(
        self, 
        doc: 'Document', 
        section: ReportSection,
        section_number: int
    ) -> None:
        """增强版的章节添加方法"""
        heading = doc.add_heading(f"{section_number}. {section.title}", level=1)
        heading.paragraph_format.space_before = Pt(24)
        
        content_para = doc.add_paragraph()
        lines = section.content.split('\n')
        for line in lines:
            if line.strip():
                run = content_para.add_run(line + '\n')
                run.font.size = Pt(11)
        
        if section.data:
            doc.add_heading("详细数据", level=2)
            
            if section.section_type == "frequency" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "")
                    data = stat.get("data", {})
                    frequencies = data.get("frequencies", [])
                    
                    doc.add_paragraph()
                    doc.add_heading(field_name, level=3)
                    
                    if frequencies:
                        table = doc.add_table(rows=1, cols=3)
                        table.style = 'Table Grid'
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        
                        hdr_cells = table.rows[0].cells
                        hdr_cells[0].text = '选项'
                        hdr_cells[1].text = '数量'
                        hdr_cells[2].text = '百分比'
                        
                        for item in frequencies[:15]:
                            row_cells = table.add_row().cells
                            row_cells[0].text = str(item.get('value', ''))
                            row_cells[1].text = str(item.get('count', 0))
                            row_cells[2].text = f"{item.get('percentage', 0)}%"
                        
                        self._format_table_headers(table)
            
            elif section.section_type == "descriptive" and "statistics" in section.data:
                for stat in section.data["statistics"]:
                    field_name = stat.get("field_name", "")
                    data = stat.get("data", {})
                    
                    doc.add_paragraph()
                    doc.add_heading(field_name, level=3)
                    
                    table = doc.add_table(rows=1, cols=8)
                    table.style = 'Table Grid'
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = '样本数'
                    hdr_cells[1].text = '均值'
                    hdr_cells[2].text = '中位数'
                    hdr_cells[3].text = '标准差'
                    hdr_cells[4].text = '最小值'
                    hdr_cells[5].text = '最大值'
                    hdr_cells[6].text = '25分位'
                    hdr_cells[7].text = '75分位'
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(data.get('count', 0))
                    row_cells[1].text = f"{data.get('mean', 0):.4f}"
                    row_cells[2].text = f"{data.get('median', 0):.4f}"
                    row_cells[3].text = f"{data.get('std', 0):.4f}"
                    row_cells[4].text = f"{data.get('min', 0):.4f}"
                    row_cells[5].text = f"{data.get('max', 0):.4f}"
                    row_cells[6].text = f"{data.get('q25', 0):.4f}"
                    row_cells[7].text = f"{data.get('q75', 0):.4f}"
                    
                    self._format_table_headers(table)
            
            elif section.section_type == "cross" and "cross_analyses" in section.data:
                for idx, analysis in enumerate(section.data["cross_analyses"]):
                    variables = analysis.get("variables", ["变量1", "变量2"])
                    cross_table = analysis.get("cross_table", [])
                    significance = analysis.get("significance")
                    
                    doc.add_paragraph()
                    doc.add_heading(f"交叉分析 {idx + 1}: {variables[0]} vs {variables[1]}", level=3)
                    
                    if cross_table:
                        first_row = cross_table[0]
                        col_values = first_row.get("col_values", {})
                        columns = [k for k in col_values.keys() if k != "_total"]
                        
                        if "count" in (col_values.get(columns[0], {}) if columns else {}):
                            table_cols = 1 + len(columns) + 1
                            table = doc.add_table(rows=1, cols=table_cols)
                            table.style = 'Table Grid'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = variables[0]
                            for i, col in enumerate(columns):
                                hdr_cells[i + 1].text = col
                            hdr_cells[-1].text = "合计"
                            
                            for row_data in cross_table:
                                row_cells = table.add_row().cells
                                row_cells[0].text = row_data.get("row", "")
                                
                                row_vals = row_data.get("col_values", {})
                                for i, col in enumerate(columns):
                                    val = row_vals.get(col, {})
                                    count = val.get("count", 0)
                                    pct = val.get("row_percentage", 0)
                                    row_cells[i + 1].text = f"{count} ({pct}%)"
                                
                                total_val = row_vals.get("_total", {})
                                row_cells[-1].text = str(total_val.get("count", 0))
                            
                            self._format_table_headers(table)
                        
                        elif "mean" in (col_values.get(columns[0], {}) if columns else {}):
                            table = doc.add_table(rows=1, cols=5)
                            table.style = 'Table Grid'
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER
                            
                            hdr_cells = table.rows[0].cells
                            hdr_cells[0].text = "组别"
                            hdr_cells[1].text = "均值"
                            hdr_cells[2].text = "标准差"
                            hdr_cells[3].text = "样本数"
                            hdr_cells[4].text = "中位数"
                            
                            for row_data in cross_table:
                                row_cells = table.add_row().cells
                                vals = row_data.get("col_values", {})
                                row_cells[0].text = row_data.get("row", "")
                                row_cells[1].text = f"{vals.get('mean', 0):.4f}"
                                row_cells[2].text = f"{vals.get('std', 0):.4f}"
                                row_cells[3].text = str(vals.get("count", 0))
                                row_cells[4].text = f"{vals.get('median', 0):.4f}"
                            
                            self._format_table_headers(table)
                    
                    self._add_significance_info_to_word(doc, significance)
        
        doc.add_paragraph()
    
    def create_export_package(
        self,
        report_id: str,
        format: str = "docx"
    ) -> Optional[str]:
        """
        创建导出包，包含Word文档和配套的Excel图表文件（压缩为ZIP）
        
        Args:
            report_id: 报告ID
            format: 导出格式（docx或pdf）
            
        Returns:
            导出文件路径（可能是ZIP包路径）
        """
        if format == "pdf":
            return self.export_to_pdf(report_id)
        
        word_path, excel_files = self.export_to_word_with_charts(
            report_id, 
            include_excel_charts=True
        )
        
        if not word_path:
            return None
        
        if not excel_files:
            return word_path
        
        zip_path = os.path.join(self.export_folder, f"{report_id}_完整导出包.zip")
        
        with ZipFile(zip_path, 'w') as zipf:
            zipf.write(word_path, os.path.basename(word_path))
            
            for excel_file in excel_files:
                if os.path.exists(excel_file):
                    arcname = f"图表附件/{os.path.basename(excel_file)}"
                    zipf.write(excel_file, arcname)
            
            readme_content = f"""报告导出包说明
================

报告文件：{os.path.basename(word_path)}
生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

文件说明：
1. Word文档：主报告文件，包含详细的数据分析结果
2. Excel图表文件（在"图表附件"文件夹中）：
   - 每个分析结果对应一个独立的Excel文件
   - Excel文件包含原始数据和可编辑的图表
   - 双击图表可在Excel中直接编辑数据和样式

使用提示：
- Word文档中的数据表格与Excel文件中的数据完全对应
- Excel图表支持修改颜色、样式、数据范围等
- 如需更新图表数据，只需修改Excel中的数据表格

"""
            zipf.writestr("使用说明.txt", readme_content.encode('utf-8'))
        
        try:
            os.remove(word_path)
        except:
            pass
        
        return zip_path
