import os
import io
import tempfile
from typing import List, Optional, Tuple, Dict, Any
from abc import ABC, abstractmethod
from PIL import Image
import fitz
from docx import Document

from app.core.config import get_settings
from app.core.logging_config import get_logger
from app.schemas.common import BoundingBox, TextBlock, ImageRegion, TableData
from app.schemas.document import PageInfo, StandardizedDocument, DocumentTypeEnum
from app.ml.ocr_engine import OCREngine

logger = get_logger(__name__)
settings = get_settings()


class BaseParser(ABC):
    def __init__(self, ocr_engine: Optional[OCREngine] = None):
        self.ocr_engine = ocr_engine or OCREngine()

    @abstractmethod
    def parse(self, file_path: str, original_filename: str) -> StandardizedDocument:
        pass

    def _run_ocr_if_needed(self, page_info: PageInfo, image: Image.Image, page_number: int) -> PageInfo:
        if not page_info.text_blocks and self.ocr_engine.is_available():
            logger.info(f"Running OCR on page {page_number} as no text was found")
            page_info.text_blocks = self.ocr_engine.ocr_image(image, page_number)
            if page_info.text_blocks:
                confidences = [tb.confidence for tb in page_info.text_blocks if tb.confidence]
                page_info.ocr_confidence = sum(confidences) / len(confidences) if confidences else 0.0
        return page_info

    def _convert_pdf_page_to_image(self, pdf_page, dpi: int = 200) -> Image.Image:
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = pdf_page.get_pixmap(matrix=mat, alpha=False)
        img_data = pix.tobytes("png")
        return Image.open(io.BytesIO(img_data)).convert("RGB")


class PDFParser(BaseParser):
    def parse(self, file_path: str, original_filename: str) -> StandardizedDocument:
        logger.info(f"Parsing PDF: {file_path}")
        import time
        start_time = time.time()

        pages: List[PageInfo] = []

        try:
            doc = fitz.open(file_path)
            page_count = len(doc)

            for page_num in range(page_count):
                page = doc[page_num]
                page_info = self._parse_page(page, page_num + 1)
                pages.append(page_info)

            doc.close()

            standardized = StandardizedDocument(
                original_filename=original_filename,
                document_type=DocumentTypeEnum.PDF,
                page_count=page_count,
                pages=pages,
                preprocessing_time=time.time() - start_time,
                metadata={"parser": "pymupdf", "page_count": page_count},
            )

            logger.info(f"PDF parsing complete: {page_count} pages")
            return standardized

        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise

    def _parse_page(self, page, page_number: int) -> PageInfo:
        rect = page.rect
        page_info = PageInfo(
            page_number=page_number,
            width=rect.width,
            height=rect.height,
        )

        text_blocks = page.get_text("blocks")
        for block in text_blocks:
            if block[6] == 0:
                x1, y1, x2, y2, text, block_type, block_no = block
                text = text.strip()
                if text:
                    bbox = BoundingBox(x1=x1, y1=y1, x2=x2, y2=y2)
                    text_block = TextBlock(
                        text=text,
                        bbox=bbox,
                        block_type="text",
                        page_number=page_number,
                        confidence=1.0,
                    )
                    page_info.text_blocks.append(text_block)

        images = page.get_images(full=True)
        for img_idx, img in enumerate(images):
            xref = img[0]
            try:
                img_rects = page.get_image_rects(xref)
                for img_rect in img_rects:
                    bbox = BoundingBox(
                        x1=img_rect.x0,
                        y1=img_rect.y0,
                        x2=img_rect.x1,
                        y2=img_rect.y1,
                    )
                    image_region = ImageRegion(
                        region_id=f"img_{page_number}_{img_idx}",
                        bbox=bbox,
                        image_type="embedded",
                        page_number=page_number,
                    )
                    page_info.image_regions.append(image_region)
            except Exception as e:
                logger.debug(f"Error extracting image region: {e}")

        try:
            tables = page.find_tables()
            for table_idx, table in enumerate(tables.tables):
                try:
                    table_data = TableData(
                        headers=table.header.names if table.header else [],
                        rows=table.extract(),
                        row_count=len(table.rows),
                        col_count=len(table.cols),
                    )
                    page_info.tables.append(table_data)
                except Exception as e:
                    logger.debug(f"Error extracting table {table_idx}: {e}")
        except Exception as e:
            logger.debug(f"Table extraction failed for page {page_number}: {e}")

        if not page_info.text_blocks:
            page_image = self._convert_pdf_page_to_image(page)
            page_info = self._run_ocr_if_needed(page_info, page_image, page_number)

        return page_info


class WordParser(BaseParser):
    def parse(self, file_path: str, original_filename: str) -> StandardizedDocument:
        logger.info(f"Parsing Word document: {file_path}")
        import time
        start_time = time.time()

        pages: List[PageInfo] = []

        try:
            doc = Document(file_path)

            page_info = PageInfo(
                page_number=1,
                width=612,
                height=792,
            )

            current_y = 50
            line_height = 12

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style_name = para.style.name if para.style else "Normal"
                    block_type = "heading" if style_name.startswith("Heading") else "text"

                    bbox = BoundingBox(
                        x1=50,
                        y1=current_y,
                        x2=562,
                        y2=current_y + line_height,
                    )

                    text_block = TextBlock(
                        text=text,
                        bbox=bbox,
                        block_type=block_type,
                        page_number=1,
                        confidence=1.0,
                    )
                    page_info.text_blocks.append(text_block)
                    current_y += line_height

            for table_idx, table in enumerate(doc.tables):
                try:
                    headers = []
                    rows = []

                    if len(table.rows) > 0:
                        first_row = table.rows[0]
                        headers = [cell.text.strip() for cell in first_row.cells]
                        rows = []
                        for row in table.rows[1:]:
                            rows.append([cell.text.strip() for cell in row.cells])

                    table_data = TableData(
                        headers=headers,
                        rows=rows,
                        row_count=len(table.rows),
                        col_count=len(table.columns) if table.rows else 0,
                    )
                    page_info.tables.append(table_data)
                except Exception as e:
                    logger.debug(f"Error extracting Word table {table_idx}: {e}")

            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        bbox = BoundingBox(
                            x1=50,
                            y1=current_y,
                            x2=550,
                            y2=current_y + 200,
                        )
                        image_region = ImageRegion(
                            region_id=f"word_img_{len(page_info.image_regions)}",
                            bbox=bbox,
                            image_type="embedded",
                            page_number=1,
                        )
                        page_info.image_regions.append(image_region)
                        current_y += 210
                    except Exception as e:
                        logger.debug(f"Error extracting Word image: {e}")

            pages.append(page_info)

            page_count = 1
            try:
                page_count = int(doc.core_properties.pages) or 1
            except Exception:
                pass

            standardized = StandardizedDocument(
                original_filename=original_filename,
                document_type=DocumentTypeEnum.WORD,
                page_count=page_count,
                pages=pages,
                preprocessing_time=time.time() - start_time,
                metadata={"parser": "python-docx"},
            )

            logger.info(f"Word parsing complete")
            return standardized

        except Exception as e:
            logger.error(f"Word parsing failed: {e}")
            raise


class ImageParser(BaseParser):
    SUPPORTED_FORMATS = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp"}

    def parse(self, file_path: str, original_filename: str) -> StandardizedDocument:
        logger.info(f"Parsing image: {file_path}")
        import time
        start_time = time.time()

        pages: List[PageInfo] = []

        try:
            image = Image.open(file_path).convert("RGB")
            width, height = image.size

            page_info = PageInfo(
                page_number=1,
                width=float(width),
                height=float(height),
            )

            page_info = self._run_ocr_if_needed(page_info, image, 1)

            pages.append(page_info)

            standardized = StandardizedDocument(
                original_filename=original_filename,
                document_type=DocumentTypeEnum.IMAGE,
                page_count=1,
                pages=pages,
                preprocessing_time=time.time() - start_time,
                metadata={"parser": "pillow", "format": image.format, "mode": image.mode},
            )

            logger.info(f"Image parsing complete")
            return standardized

        except Exception as e:
            logger.error(f"Image parsing failed: {e}")
            raise


class TXTParser(BaseParser):
    def parse(self, file_path: str, original_filename: str) -> StandardizedDocument:
        logger.info(f"Parsing TXT: {file_path}")
        import time
        start_time = time.time()

        pages: List[PageInfo] = []

        try:
            encoding = self._detect_encoding(file_path)

            with open(file_path, "r", encoding=encoding, errors="ignore") as f:
                content = f.read()

            lines = content.split("\n")
            chars_per_page = 3000
            total_chars = len(content)
            page_count = max(1, (total_chars + chars_per_page - 1) // chars_per_page)

            char_idx = 0
            for page_num in range(page_count):
                page_info = PageInfo(
                    page_number=page_num + 1,
                    width=612,
                    height=792,
                )

                current_y = 50
                line_height = 12

                chars_collected = 0
                line_idx = 0

                while char_idx < len(lines) and chars_collected < chars_per_page and line_idx < 60:
                    line = lines[char_idx]
                    char_idx += 1

                    if line.strip():
                        bbox = BoundingBox(
                            x1=50,
                            y1=current_y,
                            x2=562,
                            y2=current_y + line_height,
                        )

                        text_block = TextBlock(
                            text=line,
                            bbox=bbox,
                            block_type="text",
                            page_number=page_num + 1,
                            line_number=line_idx,
                            confidence=1.0,
                        )
                        page_info.text_blocks.append(text_block)

                        chars_collected += len(line)
                        current_y += line_height
                        line_idx += 1

                pages.append(page_info)

            standardized = StandardizedDocument(
                original_filename=original_filename,
                document_type=DocumentTypeEnum.TXT,
                page_count=page_count,
                pages=pages,
                preprocessing_time=time.time() - start_time,
                metadata={"parser": "python-builtin", "encoding": encoding},
            )

            logger.info(f"TXT parsing complete: {page_count} pages")
            return standardized

        except Exception as e:
            logger.error(f"TXT parsing failed: {e}")
            raise

    def _detect_encoding(self, file_path: str) -> str:
        encodings = ["utf-8", "gbk", "gb2312", "utf-16", "latin-1"]
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    f.read()
                return encoding
            except UnicodeDecodeError:
                continue
        return "utf-8"


class ParserFactory:
    _parsers: Dict[DocumentTypeEnum, type] = {
        DocumentTypeEnum.PDF: PDFParser,
        DocumentTypeEnum.WORD: WordParser,
        DocumentTypeEnum.IMAGE: ImageParser,
        DocumentTypeEnum.TXT: TXTParser,
    }

    @classmethod
    def get_parser(cls, doc_type: DocumentTypeEnum, ocr_engine: Optional[OCREngine] = None) -> BaseParser:
        parser_class = cls._parsers.get(doc_type)
        if not parser_class:
            raise ValueError(f"No parser available for document type: {doc_type}")
        return parser_class(ocr_engine=ocr_engine)

    @classmethod
    def detect_document_type(cls, filename: str, mime_type: Optional[str] = None) -> DocumentTypeEnum:
        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            return DocumentTypeEnum.PDF
        elif ext in {".doc", ".docx"}:
            return DocumentTypeEnum.WORD
        elif ext in ImageParser.SUPPORTED_FORMATS:
            return DocumentTypeEnum.IMAGE
        elif ext == ".txt":
            return DocumentTypeEnum.TXT

        if mime_type:
            if "pdf" in mime_type.lower():
                return DocumentTypeEnum.PDF
            elif "word" in mime_type.lower() or "officedocument.word" in mime_type.lower():
                return DocumentTypeEnum.WORD
            elif "image" in mime_type.lower():
                return DocumentTypeEnum.IMAGE
            elif "text/plain" in mime_type.lower():
                return DocumentTypeEnum.TXT

        return DocumentTypeEnum.UNKNOWN
