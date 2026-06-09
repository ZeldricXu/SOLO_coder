import os
import re
import logging
from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from pathlib import Path

from app.core.config import get_settings
from app.core.logging_config import get_logger

logger = get_logger(__name__)
settings = get_settings()


@dataclass
class DocumentTypeDetectionResult:
    has_text_layer: bool = False
    text_layer_quality: float = 0.0
    has_tables: bool = False
    table_count: int = 0
    has_handwriting: bool = False
    handwriting_confidence: float = 0.0
    is_scanned_document: bool = False
    optimal_processing_path: str = "standard"
    text_density: float = 0.0
    image_density: float = 0.0
    recommended_ocr: bool = False
    page_count: int = 0

    def to_dict(self) -> Dict[str, Any]:
        return {
            "has_text_layer": self.has_text_layer,
            "text_layer_quality": self.text_layer_quality,
            "has_tables": self.has_tables,
            "table_count": self.table_count,
            "has_handwriting": self.has_handwriting,
            "handwriting_confidence": self.handwriting_confidence,
            "is_scanned_document": self.is_scanned_document,
            "optimal_processing_path": self.optimal_processing_path,
            "text_density": self.text_density,
            "image_density": self.image_density,
            "recommended_ocr": self.recommended_ocr,
            "page_count": self.page_count,
        }


class DocumentTypeDetector:
    _instance = None

    def __new__(cls):
        if cls._instance is None:
            cls._instance = super().__new__(cls)
            cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        if self._initialized:
            return
        self._initialized = True
        self._pdf_text_cache: Dict[str, str] = {}
        self._table_keywords = [
            "表", "表格", "Table", "TABLE",
            "序号", "项目", "金额", "费用",
            "日期", "时间", "数量", "单价",
            "合计", "总计", "小计", "共计",
        ]
        self._handwriting_indicators = [
            "手写", "签字", "签名", "签章",
            "Handwritten", "Signature",
            "草书", "潦草", "涂改",
        ]

    def detect_document_type(
        self,
        file_path: str,
        original_filename: str,
        mime_type: Optional[str] = None,
    ) -> DocumentTypeDetectionResult:
        result = DocumentTypeDetectionResult()

        ext = Path(original_filename).suffix.lower()

        if ext == ".pdf":
            result.page_count = self._get_pdf_page_count(file_path)
            pdf_text = self._extract_pdf_text(file_path)

            text_char_count = len(pdf_text.strip())
            result.text_density = text_char_count / max(result.page_count, 1)

            if text_char_count > 30:
                result.has_text_layer = True
                result.text_layer_quality = min(text_char_count / 5000.0, 1.0)
                result.is_scanned_document = False
            else:
                result.has_text_layer = False
                result.is_scanned_document = True

            result.has_tables = self._detect_tables_in_pdf(file_path, pdf_text)
            if result.has_tables:
                result.table_count = self._count_tables_in_text(pdf_text)

            result.has_handwriting = self._detect_handwriting(pdf_text)
            if result.has_handwriting:
                result.handwriting_confidence = 0.7

            image_density = self._get_pdf_image_density(file_path)
            result.image_density = image_density

            if not result.has_text_layer or result.image_density > 0.7 or result.has_handwriting:
                result.recommended_ocr = True

            result.optimal_processing_path = self._determine_optimal_path(result)

        elif ext in [".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif"]:
            result.is_scanned_document = True
            result.has_text_layer = False
            result.recommended_ocr = True
            result.page_count = 1
            result.optimal_processing_path = "ocr_only"

            has_tables, has_handwriting = self._analyze_image(file_path)
            result.has_tables = has_tables
            result.has_handwriting = has_handwriting
            if has_tables:
                result.table_count = 1
            if has_handwriting:
                result.handwriting_confidence = 0.6

        elif ext in [".doc", ".docx"]:
            result.has_text_layer = True
            result.text_layer_quality = 0.9
            result.is_scanned_document = False
            result.recommended_ocr = False
            result.page_count = 1
            result.optimal_processing_path = "text_only"

            doc_text = self._extract_word_text(file_path)
            result.has_tables = self._detect_tables_in_text(doc_text)
            if result.has_tables:
                result.table_count = self._count_tables_in_text(doc_text)
            result.has_handwriting = False

        elif ext == ".txt":
            result.has_text_layer = True
            result.text_layer_quality = 1.0
            result.is_scanned_document = False
            result.recommended_ocr = False
            result.page_count = 1
            result.optimal_processing_path = "text_only"
            result.has_tables = False
            result.has_handwriting = False

        else:
            result.is_scanned_document = True
            result.recommended_ocr = True
            result.optimal_processing_path = "ocr_only"

        logger.info(
            f"Document type detection for {original_filename}: "
            f"has_text_layer={result.has_text_layer}, "
            f"is_scanned={result.is_scanned_document}, "
            f"has_tables={result.has_tables}, "
            f"has_handwriting={result.has_handwriting}, "
            f"optimal_path={result.optimal_processing_path}, "
            f"recommended_ocr={result.recommended_ocr}"
        )

        return result

    def _extract_pdf_text(self, file_path: str) -> str:
        if file_path in self._pdf_text_cache:
            return self._pdf_text_cache[file_path]

        text = ""
        try:
            import fitz

            doc = fitz.open(file_path)
            for page in doc:
                page_text = page.get_text()
                text += page_text + "\n"
            doc.close()
        except ImportError:
            logger.warning("PyMuPDF not available, using fallback text extraction")
            text = self._fallback_pdf_text_extraction(file_path)
        except Exception as e:
            logger.warning(f"Failed to extract PDF text with PyMuPDF: {e}")
            text = self._fallback_pdf_text_extraction(file_path)

        self._pdf_text_cache[file_path] = text
        return text

    def _fallback_pdf_text_extraction(self, file_path: str) -> str:
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            logger.warning(f"Fallback PDF text extraction failed: {e}")
            return ""

    def _get_pdf_page_count(self, file_path: str) -> int:
        try:
            import fitz

            doc = fitz.open(file_path)
            page_count = len(doc)
            doc.close()
            return page_count
        except Exception as e:
            logger.debug(f"Failed to get PDF page count: {e}")
            return 1

    def _get_pdf_image_density(self, file_path: str) -> float:
        try:
            import fitz

            doc = fitz.open(file_path)
            total_image_count = 0
            total_text_chars = 0

            for page in doc:
                image_list = page.get_images()
                total_image_count += len(image_list)
                page_text = page.get_text()
                total_text_chars += len(page_text.strip())

            doc.close()

            if total_image_count == 0:
                return 0.0

            if total_text_chars < 100 and total_image_count > 0:
                return 1.0

            ratio = total_image_count / max(total_text_chars / 1000, 1)
            return min(ratio, 1.0)

        except Exception as e:
            logger.debug(f"Failed to get PDF image density: {e}")
            return 0.0

    def _detect_tables_in_pdf(self, file_path: str, pdf_text: str) -> bool:
        if self._detect_tables_in_text(pdf_text):
            return True

        try:
            import fitz

            doc = fitz.open(file_path)
            for page in doc:
                drawings = page.get_drawings()
                vertical_lines = 0
                horizontal_lines = 0

                for drawing in drawings:
                    if drawing.get("type") == "l":
                        rect = drawing.get("rect")
                        if rect:
                            width = rect.width if hasattr(rect, "width") else abs(rect[2] - rect[0])
                            height = rect.height if hasattr(rect, "height") else abs(rect[3] - rect[1])

                            if width > height * 3:
                                horizontal_lines += 1
                            elif height > width * 3:
                                vertical_lines += 1

                if vertical_lines >= 2 and horizontal_lines >= 2:
                    doc.close()
                    return True

            doc.close()
        except Exception as e:
            logger.debug(f"Failed to detect tables via PDF drawings: {e}")

        return False

    def _detect_tables_in_text(self, text: str) -> bool:
        if not text:
            return False

        keyword_matches = 0
        for keyword in self._table_keywords:
            if keyword in text:
                keyword_matches += 1

        if keyword_matches >= 3:
            return True

        pipe_pattern = re.compile(r"^\s*[^|]*\|[^|]*\|")
        line_count = 0
        for line in text.split("\n"):
            if pipe_pattern.match(line):
                line_count += 1
                if line_count >= 2:
                    return True

        return False

    def _count_tables_in_text(self, text: str) -> int:
        if not text:
            return 0

        count = 0
        for keyword in self._table_keywords:
            count += text.count(keyword)

        pipe_pattern = re.compile(r"^\s*[^|]*\|[^|]*\|")
        pipe_groups = 0
        in_table = False
        for line in text.split("\n"):
            if pipe_pattern.match(line):
                if not in_table:
                    pipe_groups += 1
                    in_table = True
            else:
                in_table = False

        return max(count, pipe_groups)

    def _detect_handwriting(self, text: str) -> bool:
        if not text:
            return False

        for indicator in self._handwriting_indicators:
            if indicator in text:
                return True

        handwriting_chars = re.findall(r"[\u4e00-\u9fa5]", text)
        if len(handwriting_chars) > 50:
            variance = self._calculate_handwriting_variance(text)
            if variance > 0.3:
                return True

        return False

    def _calculate_handwriting_variance(self, text: str) -> float:
        if not text:
            return 0.0

        chars = list(text)
        if len(chars) < 10:
            return 0.0

        char_widths = []
        for i in range(min(len(chars) - 1, 100)):
            current = chars[i]
            next_char = chars[i + 1]

            if "\u4e00" <= current <= "\u9fff":
                char_widths.append(2.0)
            elif current.isascii():
                char_widths.append(1.0)
            else:
                char_widths.append(1.5)

        if not char_widths:
            return 0.0

        mean = sum(char_widths) / len(char_widths)
        variance = sum((x - mean) ** 2 for x in char_widths) / len(char_widths)
        return min(variance / 2.0, 1.0)

    def _analyze_image(self, file_path: str) -> tuple[bool, bool]:
        has_tables = False
        has_handwriting = False

        try:
            from PIL import Image

            with Image.open(file_path) as img:
                width, height = img.size

                hist = img.histogram()
                if len(hist) >= 256:
                    brightness = sum(i * hist[i] for i in range(256)) / (width * height) / 255
                    contrast = self._calculate_contrast(hist, width, height)

                    if contrast > 0.5:
                        has_tables = True

                    edge_density = self._estimate_edge_density(img)
                    if edge_density > 0.3:
                        has_handwriting = True

        except ImportError:
            logger.warning("Pillow not available for image analysis")
        except Exception as e:
            logger.debug(f"Image analysis failed: {e}")

        return has_tables, has_handwriting

    def _calculate_contrast(self, histogram, width, height):
        total_pixels = width * height
        if total_pixels == 0:
            return 0.0

        sum_brightness = sum(i * histogram[i] for i in range(256))
        mean = sum_brightness / total_pixels

        variance = sum(((i - mean) ** 2) * histogram[i] for i in range(256)) / total_pixels
        return min(variance / (256 ** 2) * 10, 1.0)

    def _estimate_edge_density(self, img) -> float:
        try:
            import numpy as np

            img_gray = img.convert("L")
            img_array = np.array(img_gray)

            edge_count = 0
            height, width = img_array.shape

            for y in range(1, height - 1):
                for x in range(1, width - 1):
                    current = int(img_array[y, x])
                    neighbors = [
                        int(img_array[y - 1, x]),
                        int(img_array[y + 1, x]),
                        int(img_array[y, x - 1]),
                        int(img_array[y, x + 1]),
                    ]
                    diffs = [abs(current - n) for n in neighbors]
                    if max(diffs) > 30:
                        edge_count += 1

            return edge_count / (height * width) if height * width > 0 else 0.0

        except Exception as e:
            logger.debug(f"Edge density estimation failed: {e}")
            return 0.0

    def _extract_word_text(self, file_path: str) -> str:
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " | "
                    text += "\n"
            return text
        except Exception as e:
            logger.warning(f"Failed to extract Word text: {e}")
            return ""

    def _determine_optimal_path(self, result: DocumentTypeDetectionResult) -> str:
        if result.has_handwriting:
            return "ocr_with_handwriting"

        if result.has_tables and result.has_text_layer:
            return "text_with_table_recognition"

        if result.has_text_layer and result.text_layer_quality >= 0.8 and not result.is_scanned_document:
            return "text_only"

        if result.is_scanned_document or not result.has_text_layer:
            return "ocr_with_layout"

        if result.text_layer_quality >= 0.5:
            return "text_with_ocr_fallback"

        return "standard"

    def get_processing_options(self, detection_result: DocumentTypeDetectionResult) -> Dict[str, Any]:
        options = {
            "use_ocr": detection_result.recommended_ocr,
            "detect_tables": detection_result.has_tables,
            "detect_handwriting": detection_result.has_handwriting,
            "prefer_text_layer": detection_result.has_text_layer and detection_result.text_layer_quality >= 0.8,
            "processing_priority": "layout" if detection_result.has_tables else "text",
        }
        return options

    def clear_cache(self):
        self._pdf_text_cache.clear()
